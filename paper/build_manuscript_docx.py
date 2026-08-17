#!/usr/bin/env python3
"""
Render `paper/manuscript.md` to a submission-shaped .docx.

The markdown is the source of record — edit that, not the .docx, and re-run this.
Anything else means the Word file and the repository drift, which is the whole
failure mode `results/` exists to prevent.

What it does beyond a plain conversion:

  * numbers figures and tables in order of appearance, so the manuscript text can
    say "Fig. 3" without anyone maintaining a counter by hand;
  * renders image alt-text as the figure caption beneath the image;
  * treats a bold `**Table N.**` paragraph directly after a table as that table's
    caption and styles it accordingly;
  * puts Appendix A on its own landscape section, because the implementation
    diagram is 2.2:1 and unreadable at portrait text width;
  * refuses to write the file if an image referenced by the markdown is missing —
    a manuscript with a hole where a figure should be is worse than no file.

    python3 paper/build_manuscript_docx.py
"""

import os
import re
import sys

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(REPO, "paper", "manuscript.md")
OUT = os.path.join(REPO, "paper", "AdaptiShield-Manuscript.docx")

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
BODY_PT = 10.0
INK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x44, 0x4A, 0x53)

TEXT_WIDTH = Inches(6.5)          # portrait, 1" margins on Letter
LAND_WIDTH = Inches(9.0)


# ─────────────────────────────────────────────────────────── inline markup

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", re.S)


def add_runs(par, text, size=BODY_PT, color=INK, base_italic=False):
    """Render **bold**, *italic* and `code` into one paragraph."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        bold = italic = False
        font = BODY_FONT
        body = piece
        if piece.startswith("**") and piece.endswith("**"):
            bold, body = True, piece[2:-2]
        elif piece.startswith("*") and piece.endswith("*"):
            italic, body = True, piece[1:-1]
        elif piece.startswith("`") and piece.endswith("`"):
            font, body = MONO_FONT, piece[1:-1]
        run = par.add_run(body)
        run.font.name = font
        run.font.size = Pt(size if font != MONO_FONT else size - 0.5)
        run.font.bold = bold
        run.font.italic = italic or base_italic
        run.font.color.rgb = color
    return par


def para(doc, text="", size=BODY_PT, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_after=4, space_before=0, italic=False, color=INK, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.06
    if indent is not None:
        pf.left_indent = indent
    if text:
        add_runs(p, text, size=size, color=color, base_italic=italic)
    return p


# ─────────────────────────────────────────────────────────── page furniture

def page_setup(section, landscape=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Inches(11), Inches(8.5)
        section.left_margin = section.right_margin = Inches(1.0)
        section.top_margin = section.bottom_margin = Inches(0.9)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = Inches(8.5), Inches(11)
        section.left_margin = section.right_margin = Inches(1.0)
        section.top_margin = section.bottom_margin = Inches(1.0)


def add_page_numbers(section):
    """A PAGE field in the footer — Word has no API-level shortcut for this."""
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, end):
        run._r.append(el)


def shade(cell, hexcolor):
    tc = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    tc.append(el)


# ─────────────────────────────────────────────────────────── block parsing

def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_sep(line):
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def parse(md):
    """Markdown → a flat list of typed blocks. Deliberately small: the manuscript
    only uses the constructs this understands, and anything else should fail
    loudly rather than render as literal asterisks."""
    blocks = []
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1
            continue

        if s == "---":
            blocks.append(("rule", None))
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            blocks.append((f"h{len(m.group(1))}", m.group(2).strip()))
            i += 1
            continue

        m = re.match(r"^!\[(.*?)\]\((.*?)\)$", s)
        if m:
            blocks.append(("image", (m.group(1), m.group(2))))
            i += 1
            continue

        if s.startswith(">"):
            quote = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append(("quote", " ".join(quote)))
            continue

        if s.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_sep(lines[i]):
                    rows.append(split_row(lines[i]))
                i += 1
            blocks.append(("table", rows))
            continue

        m = re.match(r"^([-*])\s+(.*)$", s)
        if m:
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                item = [re.sub(r"^\s*[-*]\s+", "", lines[i]).strip()]
                i += 1
                while i < len(lines) and lines[i].startswith("  ") and lines[i].strip() \
                        and not re.match(r"^\s*[-*]\s+", lines[i]):
                    item.append(lines[i].strip())
                    i += 1
                items.append(" ".join(item))
            blocks.append(("ul", items))
            continue

        m = re.match(r"^(\d+)[.)]\s+(.*)$", s)
        if m:
            items = []
            while i < len(lines) and re.match(r"^\s*\d+[.)]\s+", lines[i]):
                item = [re.sub(r"^\s*\d+[.)]\s+", "", lines[i]).strip()]
                i += 1
                while i < len(lines) and lines[i].startswith("   ") and lines[i].strip():
                    item.append(lines[i].strip())
                    i += 1
                items.append(" ".join(item))
            blocks.append(("ol", items))
            continue

        chunk = []
        while i < len(lines) and lines[i].strip() and not re.match(
                r"^\s*(#|\||>|!\[|[-*]\s|\d+[.)]\s|---$)", lines[i]):
            chunk.append(lines[i].strip())
            i += 1
        blocks.append(("p", " ".join(chunk)))
    return blocks


# ─────────────────────────────────────────────────────────── rendering

def render_table(doc, rows, width):
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Proportional columns: a "Layer" column of single digits should not take the
    # same width as a column of prose. Weights are clipped so one long cell
    # cannot starve the rest.
    weights = []
    for ci in range(ncols):
        longest = max((len(r[ci]) for r in rows if ci < len(r)), default=8)
        weights.append(min(max(longest, 8), 55))
    total = sum(weights)
    widths = [Inches(width.inches * w / total) for w in weights]

    for ci, w in enumerate(widths):
        table.columns[ci].width = w

    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci in range(ncols):
            cell = cells[ci]
            cell.width = widths[ci]
            text = row[ci] if ci < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if ci == 0
                           else WD_ALIGN_PARAGRAPH.RIGHT if _numeric(text)
                           else WD_ALIGN_PARAGRAPH.LEFT)
            add_runs(p, text, size=8.5)
            if ri == 0:
                shade(cell, "EDF1F5")
                for run in p.runs:
                    run.font.bold = True
    return table


def fig_width(path, maxw):
    """Size a figure by capping its HEIGHT, not its width.

    Sizing by width alone gives a square scatter and a wide bar chart wildly
    different visual weight on the page. Capping height instead keeps every
    figure the same size in the reader's eye, and the cap only binds when the
    figure is tall enough for it to matter."""
    with Image.open(path) as im:
        aspect = im.size[0] / im.size[1]
    cap = Inches(6.0) if maxw > Inches(7) else Inches(3.1)
    return min(maxw, Inches(cap.inches * aspect))


def _numeric(text):
    stripped = re.sub(r"[*`]", "", text).strip()
    return bool(re.match(r"^[<>~≈]?[\d.,%/ ×⁻⁶–-]+$", stripped)) and any(
        ch.isdigit() for ch in stripped)


def build():
    with open(SRC) as fh:
        md = fh.read()
    blocks = parse(md)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(BODY_PT)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    page_setup(doc.sections[0])
    add_page_numbers(doc.sections[0])

    fig_n = [0]
    width = [TEXT_WIDTH]
    # Everything between the title and the abstract is the title block, and a
    # title block is centred. Detected by position rather than by a marker, so
    # the markdown stays plain.
    in_title_block = [False]
    pending_table_caption = [False]
    missing = []

    for kind, payload in blocks:
        if kind == "h1":
            p = para(doc, payload, size=17, align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_after=10)
            for run in p.runs:
                run.font.bold = True
            in_title_block[0] = True

        elif kind == "h2":
            if payload.startswith("Appendix"):
                # the implementation diagram is 2.2:1 — give it a landscape page
                sec = doc.add_section(WD_SECTION.NEW_PAGE)
                page_setup(sec, landscape=True)
                # footer stays linked to the previous section: adding a second
                # PAGE field to the same footer part prints the number twice
                sec.footer.is_linked_to_previous = True
                width[0] = LAND_WIDTH
            p = para(doc, payload.upper(), size=11,
                     align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=5)
            for run in p.runs:
                run.font.bold = True

        elif kind == "h3":
            p = para(doc, payload, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=8, space_after=3)
            for run in p.runs:
                run.font.bold = True
                run.font.italic = True

        elif kind == "h4":
            p = para(doc, payload, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=6, space_after=2)
            for run in p.runs:
                run.font.italic = True

        elif kind == "p":
            text = payload
            if in_title_block[0]:
                if text.startswith("**Abstract**"):
                    in_title_block[0] = False
                    para(doc, text, space_before=6)
                else:
                    para(doc, text, size=9.5 if text.startswith("¹") or
                         text.startswith("Corresponding") or
                         text.startswith("**[CONFIRM") else BODY_PT,
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3,
                         color=GREY if text.startswith("**[CONFIRM") else INK)
            elif text.startswith("**Table "):
                para(doc, text, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_after=8, color=GREY)
                pending_table_caption[0] = False
            elif text.startswith("*") and text.endswith("*") and "Artifact" in text:
                para(doc, text, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_after=6, color=GREY)
            else:
                para(doc, text)

        elif kind == "ul":
            for item in payload:
                # The bullet is a literal glyph rather than a Word list style:
                # it survives a copy into any journal template unchanged.
                p = para(doc, "", space_after=2, indent=Inches(0.25))
                p.paragraph_format.first_line_indent = Inches(-0.15)
                add_runs(p, "— " + item)

        elif kind == "ol":
            for n, item in enumerate(payload, 1):
                p = para(doc, f"{n}) {item}", space_after=2, indent=Inches(0.25))
                p.paragraph_format.first_line_indent = Inches(-0.25)

        elif kind == "quote":
            p = para(doc, payload, align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=4, space_after=6, indent=Inches(0.35))
            for run in p.runs:
                run.font.italic = True

        elif kind == "rule":
            para(doc, "", space_after=2)

        elif kind == "table":
            render_table(doc, payload, width[0])
            pending_table_caption[0] = True

        elif kind == "image":
            caption, path = payload
            full = os.path.normpath(os.path.join(REPO, "paper", path))
            if not os.path.exists(full):
                missing.append(path)
                continue
            fig_n[0] += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            p.add_run().add_picture(full, width=fig_width(full, width[0]))
            cap = para(doc, f"**Fig. {fig_n[0]}.** {caption}", size=8.5,
                       align=WD_ALIGN_PARAGRAPH.LEFT, space_after=9, color=GREY)
            cap.paragraph_format.left_indent = Inches(0.0)

    if missing:
        raise SystemExit(
            "refusing to write the manuscript — referenced figures are missing:\n  "
            + "\n  ".join(missing)
            + "\nRun paper/make_figures.py and paper/make_architecture_figure.py first."
        )

    doc.save(OUT)
    words = len(re.findall(r"\b\w+\b", md))
    print(f"wrote {OUT}")
    print(f"  {len(blocks)} blocks · {fig_n[0]} figures · ~{words:,} words of source")


if __name__ == "__main__":
    sys.exit(build())
